"""Build a visually dense, internal-engineering Word document from a JSON spec."""

from __future__ import annotations

import datetime as dt
import hashlib
import io
import random
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Inches, Pt, RGBColor, Emu, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph


NAVY = RGBColor(0x1B, 0x36, 0x5D)
STEEL = RGBColor(0x2F, 0x4A, 0x6E)
BODY = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x55, 0x55)
RULE = RGBColor(0xC5, 0xCD, 0xD6)
ROW_ALT = "EEF2F6"
HEADER_BG = "1B365D"
CALLOUT_BG = "F4F1E8"
CODE_BG = "F3F4F6"
DECISION_BG = "E7EEF5"
WARN_BG = "F8EDE8"

BODY_FONT = "Liberation Serif"
SANS = "Liberation Sans"
MONO = "Liberation Mono"


@dataclass
class Theme:
    body_font: str
    sans: str
    mono: str
    body_pt: float
    line_spacing: float
    line_twips: int
    para_after: int
    first_indent_in: float
    accent: RGBColor
    steel: RGBColor
    accent_hex: str
    header_bg: str
    row_alt: str
    left_in: float
    right_in: float
    top_in: float
    bottom_in: float
    header_layout: str
    footer_layout: str
    footer_sep: str
    title_pt: float
    heading1_pt: float
    wpp: float
    version_prefix: str


_THEME: Theme | None = None


def _hex_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def theme_for(spec: dict) -> Theme:
    """Per-file chrome. Engagement sets the palette; slug hash picks the rest.

    Keep body at 12pt and spacing in a working-notes band (tighter than 1.78)
    so 40-50 pages still hold without bringing back an exhibit mill.
    """
    slug = spec.get("slug") or "doc"
    rng = random.Random(int(hashlib.md5(slug.encode()).hexdigest()[:8], 16))
    dt = (spec.get("doc_type") or "").lower()
    if "oakridge" in slug:
        palettes = [
            ("3D4F5F", "4A5A6A", "4A5560", "E8ECF0"),
            ("4A545C", "5A646C", "3F474E", "EEEFEF"),
            ("3E4A43", "52615A", "3A4540", "E7EDE8"),
        ]
        faces = [
            ("Liberation Serif", "Liberation Sans", "Liberation Mono"),
            ("DejaVu Serif", "DejaVu Sans", "DejaVu Sans Mono"),
        ]
    elif "riverview" in slug:
        palettes = [
            ("1A5F6A", "2A6B74", "1A5F6A", "E6F0F1"),
            ("215E5A", "2F6E69", "215E5A", "E5EFEE"),
            ("164E63", "2A5F72", "164E63", "E4EEF2"),
        ]
        faces = [
            ("Noto Serif", "Noto Sans", "Noto Mono"),
            ("Liberation Serif", "Liberation Sans", "Liberation Mono"),
            ("DejaVu Serif", "Liberation Sans", "Liberation Mono"),
        ]
    else:
        palettes = [
            ("1B365D", "2F4A6E", "1B365D", "EEF2F6"),
            ("1E3A5F", "334E72", "1E3A5F", "EDF1F6"),
            ("243447", "3A4C5F", "243447", "E8ECF0"),
        ]
        faces = [
            ("Liberation Serif", "Liberation Sans", "Liberation Mono"),
            ("Noto Serif", "Liberation Sans", "Liberation Mono"),
            ("DejaVu Serif", "DejaVu Sans", "DejaVu Sans Mono"),
        ]
    # Working notes often sit in sans. Designs/memos stay serif.
    if any(k in dt for k in ("meeting", "notes", "readout", "runbook", "sop")):
        faces = faces + [
            ("Liberation Sans", "Liberation Sans", "Liberation Mono"),
        ]
    accent_hex, steel_hex, header_bg, row_alt = rng.choice(palettes)
    body_font, sans, mono = rng.choice(faces)
    spacing = rng.choice([1.58, 1.60, 1.62, 1.65, 1.68, 1.70, 1.72])
    left = rng.choice([1.05, 1.10, 1.15, 1.20])
    right = rng.choice([1.05, 1.10, 1.15, 1.20])
    top = rng.choice([1.05, 1.10, 1.15, 1.20])
    bottom = rng.choice([0.95, 1.00, 1.05, 1.10])
    indent = rng.choice([0.0, 0.12, 0.18, 0.20, 0.22])
    if body_font == "Noto Serif":
        font_factor = 0.76
    elif body_font == "DejaVu Serif":
        font_factor = 0.86
    elif "Sans" in body_font and "Serif" not in body_font:
        font_factor = 1.02
    else:
        font_factor = 1.0
    width_factor = (8.5 - left - right) / 6.2
    wpp = 310.0 * (1.78 / spacing) * width_factor * font_factor
    return Theme(
        body_font=body_font,
        sans=sans,
        mono=mono,
        body_pt=12.0,
        line_spacing=spacing,
        line_twips=int(round(spacing * 240)),
        para_after=rng.choice([10, 11, 12]),
        first_indent_in=indent,
        accent=_hex_rgb(accent_hex),
        steel=_hex_rgb(steel_hex),
        accent_hex=accent_hex,
        header_bg=header_bg,
        row_alt=row_alt,
        left_in=left,
        right_in=right,
        top_in=top,
        bottom_in=bottom,
        header_layout=rng.choice(["bar", "split", "stacked"]),
        footer_layout=rng.choice(["dots", "pipe", "short"]),
        footer_sep=rng.choice(["  ·  ", "  |  ", "  /  "]),
        title_pt=rng.choice([20, 21, 22]),
        heading1_pt=rng.choice([15, 16]),
        wpp=wpp,
        version_prefix=rng.choice(["v", "Rev ", ""]),
    )


def apply_theme(theme: Theme) -> None:
    global NAVY, STEEL, HEADER_BG, ROW_ALT, BODY_FONT, SANS, MONO, _THEME
    _THEME = theme
    NAVY = theme.accent
    STEEL = theme.steel
    HEADER_BG = theme.header_bg
    ROW_ALT = theme.row_alt
    BODY_FONT = theme.body_font
    SANS = theme.sans
    MONO = theme.mono


def issue_meta(spec: dict) -> dict:
    """Version from last revision row. Status follows the document kind, not a house stamp."""
    out = dict(spec)
    hist = out.get("revision_history") or []
    if hist:
        out["version"] = str(hist[-1][0])
    dt = (out.get("doc_type") or "").lower()
    slug = out.get("slug") or "doc"
    seed = int(hashlib.md5(slug.encode()).hexdigest()[:8], 16)
    if any(k in dt for k in ("incident", "postmortem")):
        out["status"] = ("Closed", "Final", "Recorded")[seed % 3]
    elif any(k in dt for k in ("runbook", "sop")):
        out["status"] = ("In force", "Current", "Issued")[seed % 3]
    elif any(k in dt for k in ("memo", "capacity", "recommendation")):
        out["status"] = ("For decision", "Requested", "Position")[seed % 3]
    elif any(k in dt for k in ("meeting", "notes", "readout")):
        out["status"] = ("Issued", "Circulated", "Posted")[seed % 3]
    elif any(k in dt for k in ("prd", "requirement")):
        out["status"] = ("Signed", "Baseline", "Accepted")[seed % 3]
    elif any(k in dt for k in ("test", "load", "eval", "rubric")):
        out["status"] = ("In use", "Current", "Issued")[seed % 3]
    elif "adr" in dt:
        out["status"] = ("Accepted", "In effect")[seed % 2]
    else:
        out["status"] = ("Accepted", "In effect", "Current")[seed % 3]
    return out


def _set_run_font(run, name, size_pt=None, bold=None, italic=None, color=None):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")


def _set_cell_borders(cell, color="C5CDD6", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _clear_table_borders(table: Table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")


def _keep_with_next(paragraph: Paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kwn = pPr.find(qn("w:keepNext"))
    if kwn is None:
        kwn = OxmlElement("w:keepNext")
        pPr.append(kwn)


def _spacing(paragraph: Paragraph, before=0, after=8, line=240, rule="auto"):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pPr = paragraph._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:line"), str(line))
    sp.set(qn("w:lineRule"), rule)
    sp.set(qn("w:before"), str(int(before * 20)))
    sp.set(qn("w:after"), str(int(after * 20)))


def _add_page_number(paragraph: Paragraph):
    """Insert PAGE of NUMPAGES fields."""
    def _fld(begin_text=None, instr=None, separate=False, end=False):
        run = paragraph.add_run()
        r = run._r
        fld = OxmlElement("w:fldChar")
        if begin_text is not None:
            fld.set(qn("w:fldCharType"), "begin")
            r.append(fld)
            if instr:
                ir = paragraph.add_run()
                it = OxmlElement("w:instrText")
                it.set(qn("xml:space"), "preserve")
                it.text = instr
                ir._r.append(it)
        elif separate:
            fld.set(qn("w:fldCharType"), "separate")
            r.append(fld)
        elif end:
            fld.set(qn("w:fldCharType"), "end")
            r.append(fld)
        return run

    run = paragraph.add_run()
    r = run._r
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r.append(fld1)
    ir = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    ir._r.append(it)
    sep = paragraph.add_run()
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    sep._r.append(f2)


def _add_numpages(paragraph: Paragraph):
    run = paragraph.add_run()
    r = run._r
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r.append(fld1)
    ir = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " NUMPAGES "
    ir._r.append(it)
    sep = paragraph.add_run()
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    sep._r.append(f2)


def _horizontal_line(paragraph: Paragraph, color="1B365D", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_narrow_cell_width(cell, dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")


def _paragraph_text(cell, text, font=SANS, size=9, bold=False, color=BODY, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    _spacing(p, before=0, after=0, line=230)
    run = p.add_run(text)
    _set_run_font(run, font, size, bold=bold, color=color)
    return p


def build_document(spec: dict, out_path: Path) -> Path:
    theme = theme_for(spec)
    apply_theme(theme)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(theme.left_in)
    section.right_margin = Inches(theme.right_in)
    section.top_margin = Inches(theme.top_in)
    section.bottom_margin = Inches(theme.bottom_in)
    section.header_distance = Inches(0.38 if theme.header_layout == "stacked" else 0.4)
    section.footer_distance = Inches(0.36 if theme.footer_layout == "short" else 0.4)

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(theme.body_pt)
    normal.font.color.rgb = BODY
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)

    pf = normal.paragraph_format
    pf.space_after = Pt(theme.para_after)
    pf.space_before = Pt(0)
    pf.line_spacing = theme.line_spacing

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.color.rgb = NAVY
        hs.font.name = SANS
        hs.font.bold = True
        hrPr = hs.element.get_or_add_rPr()
        rFonts = hrPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            hrPr.append(rFonts)
        rFonts.set(qn("w:ascii"), SANS)
        rFonts.set(qn("w:hAnsi"), SANS)
        if i == 1:
            hs.font.size = Pt(theme.heading1_pt)
            hs.paragraph_format.space_before = Pt(20)
            hs.paragraph_format.space_after = Pt(10)
        elif i == 2:
            hs.font.size = Pt(13)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(7)
        else:
            hs.font.size = Pt(12)
            hs.paragraph_format.space_before = Pt(11)
            hs.paragraph_format.space_after = Pt(6)

    _build_header(section, spec, theme)
    _build_footer(section, spec, theme)
    _build_title_block(doc, spec, theme)

    for block in spec.get("blocks", []):
        _render_block(doc, block)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    stamp_docx_metadata(out_path, spec)
    return out_path


def _build_header(section, spec, theme: Theme):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    org = spec.get("org", "Northstar Engineering")
    kind = spec.get("doc_type", "Internal Document")
    klass = spec.get("classification", "INTERNAL")
    title = spec.get("title", "")
    rule = theme.accent_hex

    if theme.header_layout == "split":
        _spacing(p, before=0, after=0, line=200)
        run = p.add_run(org)
        _set_run_font(run, SANS, 9, bold=True, color=NAVY)
        run = p.add_run(" " * 6)
        run = p.add_run(klass)
        _set_run_font(run, SANS, 8, bold=True, color=NAVY)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2 = header.add_paragraph()
        _spacing(p2, before=0, after=2, line=200)
        _horizontal_line(p2, color=rule, sz="10")
        run = p2.add_run(kind)
        _set_run_font(run, SANS, 8, color=STEEL)
        run = p2.add_run("  ·  ")
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p2.add_run(title)
        _set_run_font(run, SANS, 8, italic=True, color=MUTED)
        return

    if theme.header_layout == "stacked":
        _spacing(p, before=0, after=0, line=200)
        run = p.add_run(org)
        _set_run_font(run, SANS, 8, bold=True, color=NAVY)
        p2 = header.add_paragraph()
        _spacing(p2, before=0, after=2, line=200)
        _horizontal_line(p2, color=rule, sz="8")
        run = p2.add_run(title)
        _set_run_font(run, SANS, 8, italic=True, color=MUTED)
        return

    _spacing(p, before=0, after=2, line=200)
    _horizontal_line(p, color=rule, sz="18")
    run = p.add_run(org)
    _set_run_font(run, SANS, 9, bold=True, color=NAVY)
    run = p.add_run("   |   ")
    _set_run_font(run, SANS, 9, color=MUTED)
    run = p.add_run(kind)
    _set_run_font(run, SANS, 9, color=STEEL)
    run = p.add_run(" " * 8)
    run = p.add_run(klass)
    _set_run_font(run, SANS, 8, bold=True, color=NAVY)
    p2 = header.add_paragraph()
    _spacing(p2, before=0, after=0, line=200)
    run = p2.add_run(title)
    _set_run_font(run, SANS, 8, italic=True, color=MUTED)


def _build_footer(section, spec, theme: Theme):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    _spacing(p, before=4, after=0, line=200)
    _horizontal_line(p, color=theme.accent_hex, sz="12" if theme.footer_layout != "short" else "8")
    sep = theme.footer_sep
    ver = theme.version_prefix + str(spec.get("version", "1.0"))
    doc_id = spec.get("doc_id", "DOC-000")
    date = spec.get("date", "")
    klass = spec.get("classification", "INTERNAL")

    if theme.footer_layout == "short":
        run = p.add_run(doc_id)
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p.add_run(sep + date + sep + "Page ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_page_number(p)
        run = p.add_run(" of ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_numpages(p)
    elif theme.footer_layout == "pipe":
        run = p.add_run(doc_id + sep + ver + sep + spec.get("status", "") + sep)
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p.add_run("Page ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_page_number(p)
        run = p.add_run(" of ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_numpages(p)
        run = p.add_run(sep + klass)
        _set_run_font(run, SANS, 8, color=MUTED)
    else:
        run = p.add_run(doc_id)
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p.add_run(sep + ver)
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p.add_run(sep + date)
        _set_run_font(run, SANS, 8, color=MUTED)
        run = p.add_run(sep + "Page ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_page_number(p)
        run = p.add_run(" of ")
        _set_run_font(run, SANS, 8, color=MUTED)
        _add_numpages(p)
        run = p.add_run(sep + klass)
        _set_run_font(run, SANS, 8, color=MUTED)
    for run in p.runs:
        if run.font.size is None:
            _set_run_font(run, SANS, 8, color=MUTED)


def _build_title_block(doc: Document, spec: dict, theme: Theme | None = None):
    theme = theme or _THEME
    kicker = doc.add_paragraph()
    _spacing(kicker, before=0, after=2, line=200)
    run = kicker.add_run(spec.get("kicker", spec.get("doc_type", "")).upper())
    _set_run_font(run, SANS, 9, bold=True, color=NAVY)

    title = doc.add_paragraph()
    _spacing(title, before=0, after=4, line=240)
    run = title.add_run(spec.get("title", ""))
    _set_run_font(run, SANS, theme.title_pt if theme else 22, bold=True, color=NAVY)

    subtitle = spec.get("subtitle")
    if subtitle:
        sp = doc.add_paragraph()
        _spacing(sp, before=0, after=8, line=230)
        run = sp.add_run(subtitle)
        _set_run_font(run, BODY_FONT, 12, italic=True, color=STEEL)

    rule = doc.add_paragraph()
    _spacing(rule, before=0, after=10, line=80)
    _horizontal_line(rule, color=(_THEME.accent_hex if _THEME else "1B365D"), sz="16")

    meta_rows = [
        ("Document ID", spec.get("doc_id", "")),
        ("Version", str(spec.get("version", "1.0"))),
        ("Status", spec.get("status", "Draft")),
        ("Date", spec.get("date", "")),
        ("Author", spec.get("author", "Aman Kumar")),
        ("Role", spec.get("role", "Software Engineer II")),
        ("Team / engagement", spec.get("team", "")),
        ("Audience", spec.get("audience", "")),
        ("Classification", spec.get("classification", "INTERNAL")),
    ]
    if spec.get("owners"):
        meta_rows.append(("Owners", spec["owners"]))
    if spec.get("related"):
        meta_rows.append(("Related", spec["related"]))

    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.autofit = True
    table.allow_autofit = True
    _clear_table_borders(table)
    for i, (k, v) in enumerate(meta_rows):
        c0, c1 = table.rows[i].cells
        _set_cell_margins(c0, 50, 50, 70, 70)
        _set_cell_margins(c1, 50, 50, 70, 70)
        _shade_cell(c0, "EEF2F6" if i % 2 == 0 else "F7F9FB")
        _shade_cell(c1, "EEF2F6" if i % 2 == 0 else "F7F9FB")
        _set_cell_borders(c0, "D5DCE4", "4")
        _set_cell_borders(c1, "D5DCE4", "4")
        _paragraph_text(c0, k, font=SANS, size=9, bold=True, color=NAVY)
        _paragraph_text(c1, str(v), font=BODY_FONT, size=10, color=BODY)
        _set_narrow_cell_width(c0, 2200)

    spacer = doc.add_paragraph()
    _spacing(spacer, before=6, after=4, line=120)

    if spec.get("revision_history"):
        h = doc.add_paragraph()
        _spacing(h, before=8, after=6, line=240)
        run = h.add_run("Revision history")
        _set_run_font(run, SANS, 11, bold=True, color=NAVY)
        _keep_with_next(h)
        _add_table(
            doc,
            ["Ver", "Date", "Author", "Notes"],
            spec["revision_history"],
            col_widths=(900, 1600, 2000, 4500),
        )

    if spec.get("summary"):
        h = doc.add_paragraph()
        _spacing(h, before=12, after=4, line=240)
        run = h.add_run("What this is")
        _set_run_font(run, SANS, 11, bold=True, color=NAVY)
        _add_para(doc, spec["summary"])


def _add_para(doc, text, first_line=False):
    p = doc.add_paragraph()
    line = _THEME.line_twips if _THEME else 427
    after = _THEME.para_after if _THEME else 12
    _spacing(p, before=0, after=after, line=line)
    indent = _THEME.first_indent_in if _THEME else 0.2
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    run = p.add_run(text)
    _set_run_font(run, BODY_FONT, 12, color=BODY)
    return p


def _page_break_before(paragraph: Paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        pb = OxmlElement("w:pageBreakBefore")
        pPr.append(pb)


def _add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1 and str(text).startswith("Appendix"):
        _page_break_before(p)
    return p


def _add_bullets(doc, items, ordered=False):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.style = doc.styles["List Number" if ordered else "List Bullet"]
        _spacing(p, before=1, after=3, line=260)
        # style may inject a run; clear and rewrite
        if p.runs:
            p.clear()
        run = p.add_run(item)
        _set_run_font(run, BODY_FONT, 12, color=BODY)


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _shade_cell(cell, HEADER_BG)
        _set_cell_borders(cell, _THEME.accent_hex if _THEME else "1B365D", "4")
        _set_cell_margins(cell, 50, 50, 70, 70)
        _paragraph_text(cell, str(h), font=SANS, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        if col_widths and j < len(col_widths):
            _set_narrow_cell_width(cell, col_widths[j])

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            bg = "FFFFFF" if i % 2 == 0 else ROW_ALT
            _shade_cell(cell, bg)
            _set_cell_borders(cell, "C5CDD6", "4")
            _set_cell_margins(cell, 50, 50, 70, 70)
            _paragraph_text(cell, str(val), font=BODY_FONT, size=9.5, color=BODY)
            if col_widths and j < len(col_widths):
                _set_narrow_cell_width(cell, col_widths[j])

    cap_space = doc.add_paragraph()
    _spacing(cap_space, before=2, after=8, line=80)
    return table


def _add_callout(doc, label, text, kind="decision"):
    bg = {"decision": DECISION_BG, "note": CALLOUT_BG, "warn": WARN_BG}.get(kind, DECISION_BG)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, bg)
    _set_cell_borders(cell, (_THEME.accent_hex if _THEME else "1B365D") if kind != "warn" else "A15C46", "10")
    _set_cell_margins(cell, 80, 90, 120, 120)
    cell.text = ""
    p0 = cell.paragraphs[0]
    _spacing(p0, before=0, after=4, line=220)
    run = p0.add_run(label.upper())
    _set_run_font(run, SANS, 8, bold=True, color=NAVY if kind != "warn" else RGBColor(0x8A, 0x3B, 0x28))
    p1 = cell.add_paragraph()
    _spacing(p1, before=0, after=0, line=250)
    run = p1.add_run(text)
    _set_run_font(run, BODY_FONT, 10.5, color=BODY)
    sp = doc.add_paragraph()
    _spacing(sp, before=4, after=8, line=80)


def _add_code(doc, text, caption=None):
    if caption:
        p = doc.add_paragraph()
        _spacing(p, before=4, after=2, line=200)
        run = p.add_run(caption)
        _set_run_font(run, SANS, 8.5, italic=True, color=MUTED)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade_cell(cell, CODE_BG)
    _set_cell_borders(cell, "C5CDD6", "6")
    _set_cell_margins(cell, 90, 90, 120, 120)
    cell.text = ""
    p = cell.paragraphs[0]
    _spacing(p, before=0, after=0, line=220)
    # preserve line breaks
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx == 0:
            run = p.add_run(line)
            _set_run_font(run, MONO, 8.5, color=BODY)
        else:
            p2 = cell.add_paragraph()
            _spacing(p2, before=0, after=0, line=220)
            run = p2.add_run(line if line else " ")
            _set_run_font(run, MONO, 8.5, color=BODY)
    sp = doc.add_paragraph()
    _spacing(sp, before=2, after=8, line=80)


def _render_block(doc, block):
    kind = block.get("type")
    if kind == "h1":
        _add_heading(doc, block["text"], 1)
    elif kind == "h2":
        _add_heading(doc, block["text"], 2)
    elif kind == "h3":
        _add_heading(doc, block["text"], 3)
    elif kind == "p":
        _add_para(doc, block["text"])
    elif kind == "bullets":
        _add_bullets(doc, block["items"], ordered=False)
    elif kind == "steps":
        _add_bullets(doc, block["items"], ordered=True)
    elif kind == "table":
        if block.get("caption"):
            p = doc.add_paragraph()
            _spacing(p, before=4, after=3, line=200)
            run = p.add_run(block["caption"])
            _set_run_font(run, SANS, 8.5, italic=True, color=MUTED)
        _add_table(doc, block["headers"], block["rows"], block.get("col_widths"))
    elif kind == "callout":
        _add_callout(doc, block.get("label", "Note"), block["text"], block.get("kind", "decision"))
    elif kind == "code":
        _add_code(doc, block["text"], block.get("caption"))
    elif kind == "quote":
        p = doc.add_paragraph()
        _spacing(p, before=4, after=8, line=260)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(block["text"])
        _set_run_font(run, BODY_FONT, 11, italic=True, color=STEEL)
        if block.get("attrib"):
            p2 = doc.add_paragraph()
            _spacing(p2, before=0, after=8, line=220)
            p2.paragraph_format.left_indent = Inches(0.3)
            run = p2.add_run("- " + block["attrib"])
            _set_run_font(run, SANS, 9, color=MUTED)
    else:
        raise ValueError(f"Unknown block type: {kind}")


_MONTHS = {
    "January": 1, "February": 2, "March": 3, "April": 4, "May": 5, "June": 6,
    "July": 7, "August": 8, "September": 9, "October": 10, "November": 11, "December": 12,
}


def _spec_datetime(s: str, hour: int, minute: int) -> dt.datetime:
    m = re.match(r"([A-Za-z]+) (\d{1,2}), (\d{4})", s or "")
    if not m:
        d = dt.datetime(2023, 6, 1, hour, minute, 0)
    else:
        d = dt.datetime(int(m.group(3)), _MONTHS[m.group(1)], int(m.group(2)), hour, minute, 0)
    return d


def stamp_docx_metadata(path: Path, spec: dict, words: int | None = None, pages: int | None = None) -> None:
    """Replace python-docx package stamps so File > Info looks like a Word file."""
    path = Path(path)
    author = spec.get("author") or "Aman Kumar"
    title = spec.get("title") or ""
    subject = spec.get("doc_type") or ""
    company = spec.get("org") or "Northstar Engineering"
    seed = int(hashlib.md5(spec.get("slug", path.stem).encode()).hexdigest()[:8], 16)
    created = _spec_datetime(spec.get("date", ""), 8 + (seed % 8), (seed // 8) % 60)
    modified = created
    hist = spec.get("revision_history") or []
    if hist:
        modified = _spec_datetime(hist[-1][1], 14 + (seed % 5), (seed // 17) % 60)
        if modified < created:
            modified = created + dt.timedelta(days=2, hours=3)
    if words is None:
        words = 0
    if pages is None:
        pages = max(1, round(words / 310)) if words else 1
    chars = words * 6
    total_min = 25 + (seed % 180)

    def iso(d: dt.datetime) -> str:
        return d.strftime("%Y-%m-%dT%H:%M:%SZ")

    core = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n"
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
        'xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        f"<dc:title>{escape(title)}</dc:title>"
        f"<dc:subject>{escape(subject)}</dc:subject>"
        f"<dc:creator>{escape(author)}</dc:creator>"
        f"<cp:keywords>{escape(spec.get('doc_id') or '')}</cp:keywords>"
        "<dc:description/>"
        f"<cp:lastModifiedBy>{escape(author)}</cp:lastModifiedBy>"
        "<cp:revision>" + str(3 + (seed % 5)) + "</cp:revision>"
        f'<dcterms:created xsi:type="dcterms:W3CDTF">{iso(created)}</dcterms:created>'
        f'<dcterms:modified xsi:type="dcterms:W3CDTF">{iso(modified)}</dcterms:modified>'
        f"<cp:category>{escape(subject)}</cp:category>"
        "</cp:coreProperties>"
    )
    app = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n"
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        "<Template>Normal.dotm</Template>"
        f"<TotalTime>{total_min}</TotalTime>"
        f"<Pages>{pages}</Pages>"
        f"<Words>{words}</Words>"
        f"<Characters>{chars}</Characters>"
        "<Application>Microsoft Office Word</Application>"
        "<DocSecurity>0</DocSecurity>"
        "<Lines>0</Lines>"
        "<Paragraphs>0</Paragraphs>"
        "<ScaleCrop>false</ScaleCrop>"
        "<Company>" + escape(company) + "</Company>"
        "<LinksUpToDate>false</LinksUpToDate>"
        f"<CharactersWithSpaces>{chars + words}</CharactersWithSpaces>"
        "<SharedDoc>false</SharedDoc>"
        "<HyperlinkBase/>"
        "<HyperlinksChanged>false</HyperlinksChanged>"
        "<AppVersion>16.0000</AppVersion>"
        "</Properties>"
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "docProps/thumbnail.jpeg":
                continue
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                data = core.encode("utf-8")
            elif item.filename == "docProps/app.xml":
                data = app.encode("utf-8")
            elif item.filename == "_rels/.rels":
                data = re.sub(
                    r'<Relationship[^>]*metadata/thumbnail[^>]*/>',
                    "",
                    data.decode("utf-8"),
                ).encode("utf-8")
            elif item.filename == "[Content_Types].xml":
                data = re.sub(
                    r'<Default Extension="jpeg"[^>]*/>',
                    "",
                    data.decode("utf-8"),
                ).encode("utf-8")
            zout.writestr(item.filename, data)
    path.write_bytes(buf.getvalue())
